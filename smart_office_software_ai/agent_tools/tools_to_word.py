from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
from typing import Type, List, Union, Dict, Optional
from langchain_core.tools import BaseTool
from pydantic import BaseModel, Field
import tempfile

from utils.oss_utils import save_file


def save_to_word(content, filename=None, title=None, style='normal'):
    """
    将输入内容保存为Word文档并上传到OSS

    参数:
        content: 输入内容，可以是以下格式之一:
                 - 字符串 (纯文本)
                 - 列表的列表 (表格数据)
                 - 字典列表 (每个字典代表一个段落或表格行)
        filename: 保存的文件名 (可选)，如果不提供则自动生成
        title: 文档标题 (可选)
        style: 文档样式 ('normal'普通, 'report'报告, 'minimal'极简)

    返回:
        包含文件信息和OSS上传结果的字典
    """
    # 创建Word文档
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal'].font.size = Pt(10.5)

    # 添加标题
    if title:
        if style == 'report':
            # 报告样式的大标题
            title_para = doc.add_paragraph(title)
            title_para.style = 'Heading 1'
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()  # 空行
        elif style == 'minimal':
            # 极简样式的小标题
            title_para = doc.add_paragraph(title)
            title_para.style = 'Heading 2'
        else:
            # 普通样式
            title_para = doc.add_paragraph(title)
            title_para.style = 'Title'

    # 处理不同类型的内容输入
    if isinstance(content, str):
        # 纯文本内容
        doc.add_paragraph(content)

    elif isinstance(content, list) and len(content) > 0:
        if isinstance(content[0], dict):
            # 字典列表 - 可以处理为表格或段落
            if all(isinstance(v, (str, int, float)) for row in content for v in row.values()):
                # 如果所有值都是简单类型，创建表格
                table = doc.add_table(rows=1, cols=len(content[0].keys()))
                table.style = 'Light Grid' if style != 'minimal' else 'Table Grid'

                # 添加表头
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(content[0].keys()):
                    hdr_cells[i].text = str(key)

                # 添加数据行
                for row in content:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row.values()):
                        row_cells[i].text = str(value)
            else:
                # 复杂字典，作为段落处理
                for item in content:
                    for key, value in item.items():
                        p = doc.add_paragraph()
                        p.add_run(f"{key}: ").bold = True
                        p.add_run(str(value))
                    doc.add_paragraph()  # 段落间空行

        elif isinstance(content[0], (list, tuple)):
            # 列表的列表 - 作为表格处理
            table = doc.add_table(rows=len(content), cols=len(content[0]))
            table.style = 'Light Grid' if style != 'minimal' else 'Table Grid'

            for i, row in enumerate(content):
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)

    # 生成文件名（如果未提供）
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"document_{timestamp}.docx"
    elif not filename.endswith('.docx'):
        filename += '.docx'

    # 创建一个临时文件对象用于上传
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        # 保存Word文件到临时文件
        doc.save(tmp_file.name)
        tmp_file.seek(0)

        # 创建一个类似文件的对象，具有filename和content_type属性
        class FileWrapper:
            def __init__(self, file_obj, filename, content_type):
                self.file = file_obj
                self.filename = filename
                self.content_type = content_type
                self.stream = file_obj

            def read(self, *args, **kwargs):
                return self.file.read(*args, **kwargs)

            def seek(self, *args, **kwargs):
                return self.file.seek(*args, **kwargs)

            def close(self):
                return self.file.close()

        file_obj = FileWrapper(
            tmp_file,
            filename=filename,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        # 上传文件到OSS
        upload_result = save_file(file_obj)

        # 关闭并删除临时文件
        tmp_file.close()
        os.unlink(tmp_file.name)

        return {
            'filename': filename,
            'upload_result': upload_result
        }


class WordContentItem(BaseModel):
    pass


class SaveToWordModel(BaseModel):
    content: Union[str, List[List[Union[str, int, float]]], List[Dict[str, Union[str, int, float]]]] = Field(
        ...,
        description="要保存的内容，可以是字符串、列表的列表(表格数据)或字典列表"
    )
    filename: Optional[str] = Field(
        None,
        description="保存的文件名(可选)，如果不提供则自动生成"
    )
    title: Optional[str] = Field(
        None,
        description="文档标题(可选)"
    )
    style: str = Field(
        'normal',
        description="文档样式('normal'普通, 'report'报告, 'minimal'极简)"
    )


class WordSaveTool(BaseTool):
    name: str = "save_to_word"
    description: str = "将数据保存为Word文档并上传到OSS，支持文本、表格数据或结构化数据"
    args_schema: Type[BaseModel] = SaveToWordModel

    def _run(
            self,
            content: Union[str, List[List[Union[str, int, float]]], List[Dict[str, Union[str, int, float]]]],
            filename: Optional[str] = None,
            title: Optional[str] = None,
            style: str = 'normal'
    ) -> dict:
        """
        将输入内容保存为Word文档并上传到OSS

        参数:
            content: 输入内容，可以是以下格式之一:
                     - 字符串 (纯文本)
                     - 列表的列表 (表格数据)
                     - 字典列表 (每个字典代表一个段落或表格行)
            filename: 保存的文件名 (可选)，如果不提供则自动生成
            title: 文档标题 (可选)
            style: 文档样式 ('normal'普通, 'report'报告, 'minimal'极简)

        返回:
            包含文件信息和OSS上传结果的字典
        """
        return save_to_word(content, filename, title, style)


def get_word_tools():
    word_save_tool = WordSaveTool()
    return word_save_tool


if __name__ == '__main__':
    # 示例1: 使用纯文本
    text_content = "这是一个示例文档内容。\n这是第二行内容。"
    res = save_to_word(text_content, "示例文档.docx", title="测试文档")
    print(res)

    # 示例2: 使用表格数据
    table_data = [
        ["姓名", "年龄", "职业"],
        ["张三", 30, "工程师"],
        ["李四", 25, "设计师"]
    ]
    save_to_word(table_data, "人员表格.docx", title="员工信息", style='report')

    # 示例3: 使用字典列表
    dict_data = [
        {"标题": "项目1", "描述": "这是第一个项目的描述", "状态": "完成"},
        {"标题": "项目2", "描述": "这是第二个项目的描述", "状态": "进行中"}
    ]
    save_to_word(dict_data, "项目报告.docx", style='minimal')